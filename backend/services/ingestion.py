import os
import uuid
from typing import List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from qdrant_client.models import PointStruct
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter

from database.models import Document, ChunkMetadata
from services.embeddings import generate_embeddings
from services.vector_store import vector_store
from core.config import settings

def extract_text_from_file(file_path: str, file_type: str) -> str:
    text = ""
    if file_type == "application/pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file_type == "text/plain" or file_type == "text/markdown":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return text

async def process_document(
    db: AsyncSession, 
    dataset_id: int, 
    file_name: str, 
    file_path: str, 
    file_type: str,
    chunk_size: int = None,
    chunk_overlap: int = None,
    chunk_method: str = "recursive",
    embedding_model: str = "BAAI/bge-small-en-v1.5"
):
    chunk_size = chunk_size or settings.CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP

    # 1. Create DB Document
    db_doc = Document(
        dataset_id=dataset_id, 
        file_name=file_name, 
        file_type=file_type,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        chunk_method=chunk_method,
        embedding_model=embedding_model
    )
    db.add(db_doc)
    await db.commit()
    await db.refresh(db_doc)

    try:
        # 2. Extract Text
        full_text = extract_text_from_file(file_path, file_type)

        # 3. Chunking
        if chunk_method == "character":
            text_splitter = CharacterTextSplitter(
                separator="\n\n",
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                is_separator_regex=False,
            )
        else:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                is_separator_regex=False,
            )
        chunks = text_splitter.split_text(full_text)

        if not chunks:
            return db_doc

        # 4. Generate Embeddings
        batch_size = 100
        points = []
        chunk_metadata_list = []
        
        for i in range(0, len(chunks), batch_size):
            chunk_batch = chunks[i:i+batch_size]
            embeddings = await generate_embeddings(chunk_batch, model_name=embedding_model)
            
            for j, text_chunk in enumerate(chunk_batch):
                global_idx = i + j
                qdrant_id = str(uuid.uuid4())
                
                # DB Metadata
                chunk_meta = ChunkMetadata(
                    document_id=db_doc.id,
                    qdrant_point_id=qdrant_id,
                    chunk_index=global_idx,
                    text_content=text_chunk,
                )
                chunk_metadata_list.append(chunk_meta)
                
                # Qdrant Point
                points.append(
                    PointStruct(
                        id=qdrant_id,
                        vector=embeddings[j],
                        payload={
                            "dataset_id": dataset_id,
                            "document_id": db_doc.id,
                            "file_name": file_name,
                            "chunk_index": global_idx,
                            "text": text_chunk
                        }
                    )
                )

        # 5. Save to Postgres
        db.add_all(chunk_metadata_list)
        await db.commit()

        # 6. Save to Qdrant
        await vector_store.initialize_collection(embedding_model)
        await vector_store.upsert_points(points, embedding_model)

        return db_doc

    except Exception as e:
        # Rollback the document creation if anything fails
        await db.delete(db_doc)
        await db.commit()
        raise e
